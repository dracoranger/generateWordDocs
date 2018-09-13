import docx

def main():
    with open("pref.ini", "r") as prefs:
        #get default name
        defaultName = prefs.readline()
        #get files to extract
        files = prefs.readline().split(',')
        #get names to change
        names = []
        while prefs:
            names.append(prefs.readline())
    fullDefault = i
    capsDefault = i.split(',')[1].toUpper()+' '+i.split(',')[0].toUpper()
    lastCapsDefault = i.split(',')[0].toUpper()
    for i in names:
        full = i
        caps = i.split(',')[1].toUpper()+' '+i.split(',')[0].toUpper()
        lastCaps = i.split(',')[0].toUpper()
        for j in files:
            if j.split('.')[-1] == 'docx' || j.split('.')[-1] == 'docm' || j.split('.')[-1] == 'docb' || j.split('.')[-1] == 'doc':
                doc = docx.Document('j')
                doc.save(j.split('.')[0]+lastCaps+j.split('.')[1])
                docNew = docx.Document(j.split('.')[0]+lastCaps+j.split('.')[1])
                for paragraph in docNew.paragraphs:
                    if fullDefault in paragraph.text:
                        paragraph.text.replace(fullDefault, full)
                    if capsDefault in paragraph.text:
                        paragraph.text.replace(capsDefault, caps)
                    if lastCapsDefault in paragraph.text:
                        paragraph.text.replace(lastCapsDefault, lastCaps)
                docNew.save(j.split('.')[0]+lastCaps+j.split('.')[1])
